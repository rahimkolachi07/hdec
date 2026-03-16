"""
maintenance/rag_engine.py
─────────────────────────
Lightweight RAG engine using OpenAI embeddings + cosine similarity.
No vector DB needed — stores everything in a JSON index file.
Low cost: uses text-embedding-3-small (cheapest) + gpt-4o-mini (cheapest capable).
"""
import os, json, hashlib, math, logging
from pathlib import Path
from datetime import datetime

logger = logging.getLogger(__name__)

# ── Index file path ────────────────────────────────────────────────────────────
def _index_path():
    from django.conf import settings
    return Path(settings.HDEC_DOCS_DIR) / "_rag_index.json"

def _docs_dir():
    from django.conf import settings
    return Path(settings.HDEC_DOCS_DIR)

def _load_index():
    p = _index_path()
    if p.exists():
        try:
            return json.loads(p.read_text(encoding='utf-8'))
        except Exception:
            pass
    return {"documents": {}, "chunks": []}

def _save_index(idx):
    _index_path().parent.mkdir(parents=True, exist_ok=True)
    _index_path().write_text(json.dumps(idx, ensure_ascii=False, indent=2), encoding='utf-8')

# ── Text extraction ────────────────────────────────────────────────────────────
def extract_text(filepath: str) -> str:
    fp = Path(filepath)
    ext = fp.suffix.lower()
    text = ""
    try:
        if ext == ".pdf":
            import PyPDF2
            with open(fp, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += (page.extract_text() or "") + "\n"
        elif ext in (".docx", ".doc"):
            import docx
            doc = docx.Document(fp)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    text += " | ".join(c.text for c in row.cells) + "\n"
        elif ext in (".txt", ".md", ".csv"):
            import chardet
            raw = fp.read_bytes()
            enc = chardet.detect(raw).get('encoding') or 'utf-8'
            text = raw.decode(enc, errors='replace')
        elif ext in (".xlsx", ".xls"):
            import pandas as pd
            xl = pd.read_excel(fp, sheet_name=None)
            for sheet_name, df in xl.items():
                text += f"\n=== {sheet_name} ===\n"
                text += df.to_string(index=False) + "\n"
        else:
            text = fp.read_text(encoding='utf-8', errors='replace')
    except Exception as e:
        logger.error(f"extract_text error {fp}: {e}")
    return text.strip()

# ── Chunking ───────────────────────────────────────────────────────────────────
def chunk_text(text: str, chunk_size: int = 600, overlap: int = 100):
    """Split text into overlapping chunks by words."""
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = " ".join(words[i:i+chunk_size])
        if chunk.strip():
            chunks.append(chunk)
        i += chunk_size - overlap
    return chunks

# ── Embeddings ─────────────────────────────────────────────────────────────────
def get_embedding(text: str, api_key: str) -> list:
    """Get embedding using text-embedding-3-small (cheapest)."""
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    resp = client.embeddings.create(
        model="text-embedding-3-small",
        input=text[:8000]
    )
    return resp.data[0].embedding

def cosine_similarity(a: list, b: list) -> float:
    dot = sum(x*y for x,y in zip(a,b))
    na  = math.sqrt(sum(x*x for x in a))
    nb  = math.sqrt(sum(x*x for x in b))
    if na == 0 or nb == 0:
        return 0.0
    return dot / (na * nb)

# ── Index a document ──────────────────────────────────────────────────────────
def index_document(filepath: str, filename: str, api_key: str) -> dict:
    """
    Extract text, chunk, optionally embed, save to index.
    If OpenAI API is unreachable (e.g. PythonAnywhere free tier),
    chunks are saved WITHOUT embeddings — keyword search still works.
    """
    # Step 1: Extract text
    try:
        text = extract_text(filepath)
    except Exception as e:
        return {"error": "Text extraction error: " + str(e)}

    if not text or not text.strip():
        return {"error": "Could not extract text — file may be empty, password-protected, or a scanned image"}

    # Step 2: Chunk
    chunks = chunk_text(text)
    if not chunks:
        return {"error": "Document appears to be empty after chunking"}

    doc_id = hashlib.md5(filename.encode()).hexdigest()[:12]
    idx = _load_index()
    idx["chunks"] = [c for c in idx["chunks"] if c.get("doc_id") != doc_id]

    # Step 3: Embed (optional — graceful fallback if OpenAI unreachable)
    embedded_chunks = []
    embed_error = ""
    embed_ok = False

    if api_key:
        try:
            import urllib.request, urllib.error
            # Test connectivity first with a tiny request
            test_req = urllib.request.Request(
                "https://api.openai.com/v1/models",
                headers={"Authorization": "Bearer " + api_key},
                method="GET"
            )
            urllib.request.urlopen(test_req, timeout=8)

            # Connectivity OK — embed using urllib (no openai package needed)
            for batch_start in range(0, len(chunks), 20):
                batch = chunks[batch_start:batch_start+20]
                payload = json.dumps({
                    "model": "text-embedding-3-small",
                    "input": [c[:8000] for c in batch]
                }).encode("utf-8")
                req = urllib.request.Request(
                    "https://api.openai.com/v1/embeddings",
                    data=payload,
                    headers={"Authorization": "Bearer " + api_key, "Content-Type": "application/json"},
                    method="POST"
                )
                with urllib.request.urlopen(req, timeout=60) as resp:
                    result = json.loads(resp.read().decode("utf-8"))
                for j, emb_data in enumerate(result["data"]):
                    embedded_chunks.append({
                        "doc_id":    doc_id,
                        "filename":  filename,
                        "chunk_idx": batch_start + j,
                        "text":      batch[j],
                        "embedding": emb_data["embedding"],
                    })
            embed_ok = True
        except Exception as e:
            embed_error = str(e)
            logger.warning(f"[RAG] Embedding skipped (no internet or API error): {e}")

    # If embedding failed or skipped — store chunks without embeddings
    # (keyword search will still work in rag_chat)
    if not embed_ok:
        for i, chunk in enumerate(chunks):
            embedded_chunks.append({
                "doc_id":    doc_id,
                "filename":  filename,
                "chunk_idx": i,
                "text":      chunk,
                "embedding": [],   # empty = keyword-only search
            })

    # Step 4: Save
    idx["documents"][doc_id] = {
        "filename":   filename,
        "filepath":   filepath,
        "indexed_at": datetime.now().strftime("%d %b %Y %H:%M"),
        "chunks":     len(embedded_chunks),
        "chars":      len(text),
        "doc_id":     doc_id,
        "embedded":   embed_ok,
    }
    idx["chunks"].extend(embedded_chunks)
    _save_index(idx)

    status_msg = "embedded" if embed_ok else ("stored (no embeddings: " + embed_error[:80] + ")" if embed_error else "stored without embeddings")
    return {
        "doc_id":  doc_id,
        "chunks":  len(embedded_chunks),
        "chars":   len(text),
        "success": True,
        "embedded": embed_ok,
        "status": status_msg,
    }

# ── Remove a document ─────────────────────────────────────────────────────────
def remove_document(doc_id: str):
    idx = _load_index()
    doc = idx["documents"].pop(doc_id, None)
    idx["chunks"] = [c for c in idx["chunks"] if c.get("doc_id") != doc_id]
    _save_index(idx)
    # Delete file
    if doc:
        fp = Path(doc.get("filepath",""))
        if fp.exists():
            try: fp.unlink()
            except: pass
    return bool(doc)

# ── Search ────────────────────────────────────────────────────────────────────
def search_chunks(query: str, api_key: str, top_k: int = 5) -> list:
    """
    Search chunks by embedding similarity (if available) or keyword fallback.
    Works on PythonAnywhere free tier via keyword matching.
    """
    idx = _load_index()
    if not idx["chunks"]:
        return []

    # Check if any chunks have embeddings
    has_embeddings = any(c.get("embedding") for c in idx["chunks"])

    if has_embeddings and api_key:
        # Try vector search first
        try:
            import urllib.request
            payload = json.dumps({
                "model": "text-embedding-3-small",
                "input": query[:8000]
            }).encode("utf-8")
            req = urllib.request.Request(
                "https://api.openai.com/v1/embeddings",
                data=payload,
                headers={"Authorization": "Bearer " + api_key, "Content-Type": "application/json"},
                method="POST"
            )
            with urllib.request.urlopen(req, timeout=10) as resp:
                result = json.loads(resp.read().decode("utf-8"))
            q_emb = result["data"][0]["embedding"]
            scored = []
            for chunk in idx["chunks"]:
                emb = chunk.get("embedding")
                if emb:
                    score = cosine_similarity(q_emb, emb)
                    scored.append((score, chunk))
            scored.sort(key=lambda x: x[0], reverse=True)
            return [c for _, c in scored[:top_k]]
        except Exception as e:
            logger.warning(f"[RAG] Vector search failed, using keyword: {e}")

    # Keyword fallback — score by word overlap
    q_words = set(query.lower().split())
    scored = []
    for chunk in idx["chunks"]:
        text = chunk.get("text", "").lower()
        # Count how many query words appear in chunk
        score = sum(1 for w in q_words if len(w) > 2 and w in text)
        if score > 0:
            scored.append((score, chunk))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [c for _, c in scored[:top_k]]

# ── RAG Chat ──────────────────────────────────────────────────────────────────
def rag_chat(messages: list, plant_context: str, api_key: str) -> str:
    """
    Full RAG chat using gpt-4o-mini with plant data context + document RAG.
    """
    from openai import OpenAI
    client = OpenAI(api_key=api_key)

    # Get last user query
    query = ""
    for m in reversed(messages):
        if m.get("role") == "user":
            query = m.get("content", "")
            break

    # Retrieve relevant document chunks
    chunks = search_chunks(query, api_key, top_k=6) if query else []
    doc_context = ""
    if chunks:
        doc_context = "\n\n=== RELEVANT UPLOADED DOCUMENTS ===\n"
        seen = set()
        for c in chunks:
            fname = c.get("filename", "")
            if fname not in seen:
                doc_context += f"\n--- From: {fname} ---\n"
                seen.add(fname)
            doc_context += c.get("text", "") + "\n"

    system_prompt = f"""You are HDEC Bot — the professional AI maintenance assistant for Al Henakiya 1100 MW Solar Power Plant operated by Power China HDEC, Saudi Arabia.

You have access to LIVE PLANT DATA updated every 5 minutes from Google Sheets, plus any uploaded maintenance documents.

RESPONSE FORMAT — CRITICAL:
- Use markdown formatting: **bold** for labels, ## for section headers, - for bullet lists, 1. for numbered steps
- For status queries: give Total, Done, Pending, Rate as a clear summary then list specifics
- For fault queries: list each item with block/equipment/issue on its own line
- For compliance: show weekly breakdown table then monthly total
- Keep responses complete but concise — no waffle
- End responses about critical issues with "Recommended Action:" section
- Do NOT add extra asterisks or symbols outside of proper markdown syntax
- Use section breaks with ---  between major sections

{plant_context}
{doc_context}

IMPORTANT: The plant data above contains full details including individual records for trackers, strings, inverters, equipment failures, and observations. Use this data to give specific, accurate answers."""

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        max_tokens=1000,
        temperature=0.1,
        messages=[{"role": "system", "content": system_prompt}] + messages[-12:]
    )
    return resp.choices[0].message.content

# ── List documents ────────────────────────────────────────────────────────────
def list_documents() -> list:
    idx = _load_index()
    return sorted(idx["documents"].values(), key=lambda d: d.get("indexed_at",""), reverse=True)

# ── Stats ─────────────────────────────────────────────────────────────────────
def rag_stats() -> dict:
    idx = _load_index()
    return {
        "document_count": len(idx["documents"]),
        "chunk_count":    len(idx["chunks"]),
        "documents":      list(idx["documents"].values()),
    }
